from flask import Flask, request, send_file
from flask_cors import CORS
import yaml
from docx import Document
import os

app = Flask(__name__)
CORS(app, resources={r"/convert": {"origins": "https://yaml-to-docx-converter.vercel.app"}})

# Function to add YAML data to DOCX
def add_dict_to_docx(data, parent):
    if isinstance(data, dict):
        for key, value in data.items():
            if not isinstance(key, str):
                key = str(key)
            if isinstance(value, dict):
                parent.add_heading(key, level=1)
                add_dict_to_docx(value, parent)
            elif isinstance(value, list):
                parent.add_heading(key, level=2)
                for item in value:
                    if isinstance(item, dict):
                        add_dict_to_docx(item, parent)
                    else:
                        parent.add_paragraph(f'- {str(item)}')
            else:
                if not isinstance(value, str):
                    value = str(value)
                parent.add_paragraph(f'{key}: {value}')
    elif isinstance(data, list):
        for item in data:
            add_dict_to_docx(item, parent)
    else:
        if not isinstance(data, str):
            data = str(data)
        parent.add_paragraph(data)

# Endpoint to handle file conversion
@app.route('/convert', methods=['OPTIONS'])
def handle_options():
    return '', 200  # Respond with a 200 status for OPTIONS request

@app.route('/convert', methods=['POST'])
def convert_yaml_to_docx():
    if 'file' not in request.files:
        return 'No file part', 400

    file = request.files['file']

    if file.filename == '':
        return 'No selected file', 400

    if file and file.filename.endswith('.yaml'):
        # Save the YAML file temporarily
        file_path = os.path.join('/tmp', file.filename)
        file.save(file_path)

        # Load YAML data
        with open(file_path, 'r') as f:
            yaml_data = yaml.safe_load(f)

        # Create DOCX document
        document = Document()
        document.add_heading('YAML to DOCX', 0)

        # Add YAML data to DOCX
        add_dict_to_docx(yaml_data, document)

        # Save DOCX file temporarily
        docx_file_path = os.path.splitext(file_path)[0] + '.docx'
        document.save(docx_file_path)

        # Send DOCX file as response
        return send_file(docx_file_path, as_attachment=True, download_name=f'{file.filename.replace(".yaml", ".docx")}')

    return 'Invalid file format', 400

if __name__ == '__main__':
    app.run(debug=True)
